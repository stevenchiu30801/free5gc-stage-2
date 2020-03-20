#!/usr/bin/env python3

from docx import Document
import re

def stringToCapitalWithoutSpace(origin : str) -> str :
    ret = [s.strip() if s.isupper() else s.strip().title() for s in origin.split()]
    if ret[0][0].isdigit() == True :
        if len(ret) == 1 :
            if ret[0][0:3] == '5GS' :
                return ret[0][3:] + '5GS'
            elif ret[0][0:2] == '5G' :
                return ret[0][2:] + '5G'
            else :
                print('[stringToCapitalWithoutSpace] Special Case Need To Do')
        ret.append(ret[0])
        return ''.join(ret[1:])
    return ''.join(ret)

def stringFirstCharToLittle(origin : str) -> str :
    return origin[0].lower() + origin[1:]

def printGolangConstStruct(keyValue : dict(), fd : any) -> None :
    if len(keyValue) == 0 :
        return
    print('const (', file=fd)
    for name, val in keyValue.items() :
        print(f"    {name} uint8 = {val}", file=fd)
    print(')\n', file=fd)    

def printGolangStruct(structName : str, content : list(), comment : list, fd : any) -> None: 
    for line in comment :
        print(f'// {line}', file=fd)
    print('type', structName, 'struct {', file=fd)
    for line in content :
        print(f'    {line}', file=fd)
    print('}\n', file=fd)

def printGolangStructFunc(funcDef : dict, content : list(), fd : any) -> None :
    for comment in funcDef.get('comment', []) :
        print(f"// {comment}", file=fd)
    bindWithStruct = f" (a *{funcDef['structName']})" if funcDef.get('structName', None) != None else ''
    returnType = f" ({stringFirstCharToLittle(funcDef['returnName'])} {funcDef['returnType']})" if funcDef.get('returnType', None) != None else ''
    parameterType = f"({stringFirstCharToLittle(funcDef['parameterName'])} {funcDef['parameterType']})" if funcDef.get('parameterType', None) != None else '()'
    print(f"func{bindWithStruct}", f"{funcDef.get('funcName', '')}{parameterType}{returnType}", '{', file=fd)
    for line in content : 
        print(f'    {line}', file=fd)
    print('}\n', file=fd)

def printGolangPackage(packageName : str, fd : any) -> None :
    print(f"package {packageName}\n", file=fd)

class MsgTypeContent() :
    def __init__(self, table : list) :
        self.MessageType = list()
        self.name = ''
        self.MsgList = list()
        self.getInfo(table)

    def getInfo(self, table : list) -> None :
        titleList = list()
        for cells in table.rows[1].cells :
            titleList.append(cells.text)

        for row in table.rows[5:] :
            tmpDict = dict()
            for title, cell in zip(titleList, row.cells) :
                tmpDict[title] = re.sub('[()]', '', cell.text.strip())
            self.MessageType.append(tmpDict)

    def generateMsgType(self) -> dict :
        msgNameToValue, msgValue = dict(), 0
        for msg in self.MessageType :
            msgValue, skip = 0, False
            for key, value in msg.items() :
                if key.isdigit() == True :
                    try : 
                        msgValue += 2 ** (int(key) - 1) * int(value)
                    except : 
                        skip = True
                        break
            if skip == False :       
                msgNameToValue[f"MsgType{stringToCapitalWithoutSpace(msg['Messages'])}"] = msgValue
        return msgNameToValue

class GmmGsmContent() :
    def __init__(self, name, section, table : list) :
        self.Message = list()
        self.IE = list()
        self.name = name.strip()
        self.section = section.strip()
        self.getInfo(table)

    def getInfo(self, table : list) -> None :
        titleList = list()
        for cells in table.rows[0].cells :
            titleList.append(cells.text)

        for row in table.rows[1:] :
            # TODO : DNN now is TBD so ignore it
            tmpDict= dict()
            for title, cell in zip(titleList, row.cells) :
                #if 'DNN' in cell.text.strip() :
                #    isDNN = 1
                tmpDict[title] = cell.text.strip()
            #if isDNN == 0 :
            self.IE.append(tmpDict)

    def generateGmmGsm(self) -> list:
        ieCollection = list()
        varName, isHalf = list(), 0
        isFixed = re.compile('(\d+|[A-Za-z]+)')
        for ie in self.IE :
            varName.append(stringToCapitalWithoutSpace(re.sub('[\'\-\"]', '', ie['Information Element'])).split('(')[0])
            if ie['Length'] != '1/2' and isHalf == 1 :
                print('[Not a commplete byte]')
            if ie['Length'] == '1/2' and isHalf == 0 :
                isHalf = 1
                continue

            genDict = dict()
            varName.reverse()
            
            # print(ie['Type/Reference'].split()[-1])
            # print(ie['Information Element'])
            if ie['Type/Reference'].split()[-1] == '9.11.3.41' and ie['Information Element'] == 'PDU session ID' :
                varName = list()
                varName.append('PduSessionID2Value')

            if ie['Type/Reference'].split()[-1] == '9.11.3.41' and ie['Information Element'] == 'Old PDU session ID' :
                varName = list()
                varName.append('OldPDUSessionID')

            genDict['Name'], genDict['Presence'] = 'And'.join(varName), ie['Presence']
            genDict['Format'], genDict['Length'], genDict['Iei'] = ie['Format'], ie['Length'],  ie['IEI']
            ieCollection.append(genDict)
            varName, isHalf = list(), 0
        if isHalf == 1 :
            print('[Not a commplete byte]')
        return ieCollection

    def generateIEs(self) -> list :
        contentList, commentList, varName, isHalf = list(), list(), list(), 0
        for ie in self.IE :
            varName.append(stringToCapitalWithoutSpace(re.sub('[\'\-\"]', '', ie['Information Element'])).split('(')[0])
            commentList.append(ie['Type/Reference'].split()[-1])
            if ie['Length'] != '1/2' and isHalf == 1 :
                print('[Not a commplete byte]')
            if ie['Length'] == '1/2' and isHalf == 0 :
                isHalf = 1
                continue
            varName.reverse()         

            ieName = 'And'.join(varName)
            
            if ie['Type/Reference'].split()[-1] == '9.11.3.41' and 'OldPDUSessionID' in varName :
                ieName = 'OldPDUSessionID'  
            
            if ie['Type/Reference'].split()[-1] == '9.11.3.41' and 'PDUSessionID' in varName :
                ieName = 'PduSessionID2Value'  

            if ie['Presence'] == 'M' :
                contentList.append(f'nasType.{ieName}')
            else :
                contentList.append(f'*nasType.{ieName}')
            commentList, varName, isHalf = list(), list(), 0     
        if isHalf == 1 :
            print('[Not a commplete byte]')
        return contentList, commentList

def printGmmDecodeFunc(msgTypeKeyValue : dict(), fd : any) -> None : 
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = 'Message', 'GmmMessageDecode'
    funcDef['parameterName'], funcDef['parameterType'] = 'byteArray', '*[]byte'
    content.append('buffer := bytes.NewBuffer(*byteArray)')
    content.append('a.GmmMessage = NewGmmMessage()')
    content.append('binary.Read(buffer, binary.BigEndian, &a.GmmMessage.GmmHeader)')
    content.append('switch a.GmmMessage.GmmHeader.GetMessageType() {')
    for name, _ in msgTypeKeyValue.items() : 
        content.append(f"case {name}:")
        content.append(f"    a.GmmMessage.{name[7:]} = nasMessage.New{name[7:]}({name})")
        content.append(f"    a.GmmMessage.Decode{name[7:]}(byteArray)")
    content.append('default:')
    content.append('}')
    printGolangStructFunc(funcDef, content, fd)

def printGsmDecodeFunc(msgTypeKeyValue : dict(), fd : any) -> None : 
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = 'Message', 'GsmMessageDecode'
    funcDef['parameterName'], funcDef['parameterType'] = 'byteArray', '*[]byte'
    content.append('buffer := bytes.NewBuffer(*byteArray)')
    content.append('a.GsmMessage = NewGsmMessage()')
    content.append('binary.Read(buffer, binary.BigEndian, &a.GsmMessage.GsmHeader)')
    content.append('switch a.GsmMessage.GsmHeader.GetMessageType() {')
    for name, _ in msgTypeKeyValue.items() : 
        content.append(f"case {name}:")
        content.append(f"    a.GsmMessage.{name[7:]} = nasMessage.New{name[7:]}({name})")
        content.append(f"    a.GsmMessage.Decode{name[7:]}(byteArray)")
    content.append('default:')
    content.append('}')
    printGolangStructFunc(funcDef, content, fd)

def printGmmEncodeFunc(msgTypeKeyValue : dict(), fd : any) -> None : 
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = 'Message', 'GmmMessageEncode'
    funcDef['parameterName'], funcDef['parameterType'] = 'buffer', '*bytes.Buffer'
    content.append('switch a.GmmMessage.GmmHeader.GetMessageType() {')
    for name, _ in msgTypeKeyValue.items() :
        content.append(f"case {name}:")
        content.append(f"    a.GmmMessage.Encode{name[7:]}(buffer)")
    content.append('default:')
    content.append('}')
    printGolangStructFunc(funcDef, content, fd)

def printGsmEncodeFunc(msgTypeKeyValue : dict(), fd : any) -> None : 
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = 'Message', 'GsmMessageEncode'
    funcDef['parameterName'], funcDef['parameterType'] = 'buffer', '*bytes.Buffer'
    content.append('switch a.GsmMessage.GsmHeader.GetMessageType() {')
    for name, _ in msgTypeKeyValue.items() :
        content.append(f"case {name}:")
        content.append(f"    a.GsmMessage.Encode{name[7:]}(buffer)")
    content.append('default:')
    content.append('}')
    printGolangStructFunc(funcDef, content, fd)

def printGolangEncodeFunc(funcName : str, ieCollection : list, fd : any) -> None :
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = funcName, f"Encode{funcName}"
    funcDef['parameterName'], funcDef['parameterType'] = 'buffer', '*bytes.Buffer'
    for ie in ieCollection :
        if ie['Format'] == 'V' :
            content.append(f"binary.Write(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
        elif 'L' in ie['Format'] and 'M' in ie['Presence'] :
            content.append(f"binary.Write(buffer, binary.BigEndian, a.{ie['Name']}.GetLen())")
            content.append(f"binary.Write(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
        elif 'O' in ie['Presence'] or 'C' in ie['Presence'] :
            content.append(f"if a.{ie['Name']} != nil" + ' {')
            if 'T' in ie['Format'] and '1' != ie['Length']:
                content.append(f"    binary.Write(buffer, binary.BigEndian, a.{ie['Name']}.GetIei())")
            if 'L' in ie['Format'] :
                content.append(f"    binary.Write(buffer, binary.BigEndian, a.{ie['Name']}.GetLen())")
            if 'V' in ie['Format'] : 
                if 'Buffer' in ie['BufName'] and '1' != ie['Length']:
                    content.append(f"    binary.Write(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
                elif 'Octet' in ie['BufName'] :
                    if 'TLV' in ie['Format'] and '3' == ie['Length'] :
                        content.append(f"    binary.Write(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
                    elif 'L' in ie['Format'] :
                        content.append(f"    binary.Write(buffer, binary.BigEndian, a.{ie['Name']}.{ie['BufName']}[:a.{ie['Name']}.GetLen()])")
                    else :
                        content.append(f"    binary.Write(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
            content.append('}')
        else:
            print(f"[Error] printGolangEncodeFunc {ie['Name']} generated error")
    printGolangStructFunc(funcDef, content, fd)

def printGolangDecodeFunc(funcName : str, ieCollection : list, fd : any) -> None : 
    funcDef, content = dict(), list()
    funcDef['structName'], funcDef['funcName']  = funcName, f"Decode{funcName}"
    funcDef['parameterName'], funcDef['parameterType'] = 'byteArray', '*[]byte'
    content.append('buffer := bytes.NewBuffer(*byteArray)')
    for ie in ieCollection : 
        if ie['Format'] == 'V':
            content.append(f"binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
        elif 'L' in ie['Format'] and 'M' in ie['Presence'] :
            content.append(f"binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.Len)")
            content.append(f"a.{ie['Name']}.SetLen(a.{ie['Name']}.GetLen())")            
            content.append(f"binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
            
    content.append('for buffer.Len() > 0 {')
    content.append('    var ieiN uint8')
    content.append('    var tmpIeiN uint8')
    content.append('    binary.Read(buffer, binary.BigEndian, &ieiN)')
    content.append('    // fmt.Println(ieiN)')
    content.append('    if ieiN >= 0x80 {')
    content.append('        tmpIeiN = (ieiN & 0xf0) >> 4')
    content.append('    } else {')
    content.append('        tmpIeiN = ieiN')
    content.append('    }')
    content.append('    // fmt.Println("type", tmpIeiN)')
    content.append('    switch tmpIeiN  {')
    for ie in ieCollection :
        content.append(f"    case {funcName}{ie['Name']}Type:")
        if 'TV' in ie['Format'] and 'O' in ie['Presence'] and ie['Length'] != "1" :
            content.append(f"        a.{ie['Name']} = nasType.New{ie['Name']}(ieiN)")            
            content.append(f"        binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
        elif 'TV' in ie['Format'] and 'O' in ie['Presence'] and ie['Length'] == "1" :
            content.append(f"a.{ie['Name']} = nasType.New{ie['Name']}(ieiN)")
            content.append(f"a.{ie['Name']}.Octet = ieiN")
        elif 'TV' in ie['Format'] and 'C' in ie['Presence'] :
            content.append(f"		a.{ie['Name']} = nasType.New{ie['Name']}(ieiN)")			
            content.append(f"		binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
        elif 'TLV' in ie['Format'] and 'O' in ie['Presence'] :
            content.append(f"        a.{ie['Name']} = nasType.New{ie['Name']}(ieiN)")
            content.append(f"        binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.Len)")
            content.append(f"        a.{ie['Name']}.SetLen(a.{ie['Name']}.GetLen())")
            if ie['Length'] == '3' and len(ie['Length']) == 1:
                content.append(f"        binary.Read(buffer, binary.BigEndian, &a.{ie['Name']}.{ie['BufName']})")
            else :
                content.append(f"        binary.Read(buffer, binary.BigEndian, a.{ie['Name']}.{ie['BufName']}[:a.{ie['Name']}.GetLen()])")
        else : 
            content.pop()
    content.append(f"    default:")
    content.append('    }')
    content.append('}')
    printGolangStructFunc(funcDef, content, fd)

def generateMessageAndAPI(ieStructToBufNameDict : dict) :
    contentDoc = Document('content_type.docx')
    nasMessageNameList = list() # [nasMessageName, chapterSection]
    for line in contentDoc.paragraphs :
        if line.style.name == 'Heading 3' :
            h3TitleList = line.text.strip().split()
            nasMessageNameList.append([' '.join(h3TitleList[1:]), h3TitleList[0]])

    nasMessageList = list()
    for nasMessageName, table in zip(nasMessageNameList, contentDoc.tables) :
        nasMessageName[0] = re.sub('[()]', '', stringToCapitalWithoutSpace(re.sub('[\'\-\"]', '', nasMessageName[0])))
        nasMessageList.append(GmmGsmContent(*nasMessageName, table))
    
    for nasMessage in nasMessageList :
        fd = open(f"nasMessage/NAS_{nasMessage.name}.go", 'w')
        content, comment = nasMessage.generateIEs()
        printGolangPackage("nasMessage",fd)
        print('import (', file=fd)
        print('        "bytes"', file=fd)
        print('        "encoding/binary"', file=fd)
        # print('        "fmt"', file=fd)
        print('        "gofree5gc/lib/nas/nasType"', file=fd)
        print(')', file=fd)     
	
        printGolangStruct(nasMessage.name, content, comment, fd)

        # Like constructor
        newFuncDict, newFuncContent = dict(), list()
        newFuncDict['funcName'] = f"New{nasMessage.name}"
        newFuncDict['returnName'], newFuncDict['returnType'] = nasMessage.name, f"*{nasMessage.name}"
        newFuncDict['parameterName'] = 'Iei'
        newFuncDict['parameterType'] = 'uint8'
        newFuncContent.append(f"{stringFirstCharToLittle(nasMessage.name)} = &{nasMessage.name}" + '{}')        
        newFuncContent.append(f"return {stringFirstCharToLittle(nasMessage.name)}")
        printGolangStructFunc(newFuncDict, newFuncContent, fd)
        
        nasIeDictList, nasIeiDict = nasMessage.generateGmmGsm(), dict()
        for ie in nasIeDictList :
            ie['BufName'] = ieStructToBufNameDict.get(ie['Name'], None)
            print(f"[Debug] {ie['Name']} : {ie['BufName']}")
            if ie['BufName'] == None :
                print(f"[Error] The ie '{ie['Name']}' does not has buffer name")
            if ie.get('Iei', '') != '' :
                if ie['Iei'][-1] == '-' :
                    nasIeiDict[f"{nasMessage.name}{ie['Name']}Type"] = f"0x0{ie['Iei'][:-1]}"
                else : 
                    nasIeiDict[f"{nasMessage.name}{ie['Name']}Type"] = f"0x{ie['Iei']}"
        
        printGolangConstStruct(nasIeiDict, fd)
        printGolangEncodeFunc(nasMessage.name, nasIeDictList, fd)
        printGolangDecodeFunc(nasMessage.name, nasIeDictList, fd)
        fd.close()

    with open('nas.go', 'w') as fd_nas :
        GmmMessage, GsmMessage = list(), list()
        for nasMessageName in nasMessageNameList:
            if '8.2' in nasMessageName[1] :
                GmmMessage.append(f"*nasMessage.{nasMessageName[0]} //{nasMessageName[1]}")
            else :    
                GsmMessage.append(f"*nasMessage.{nasMessageName[0]} //{nasMessageName[1]}")    

        msgDoc = Document('MsgType_24501.docx')
        printGolangPackage("nas", fd_nas)
        printGolangStruct('Message', ['*GmmMessage', '*GsmMessage'], [], fd_nas)

        # Print GMM
        gmmMsgType = MsgTypeContent(msgDoc.tables[0])
        gmmConstKeyValue = gmmMsgType.generateMsgType()
        printGolangStruct('GmmMessage', GmmMessage, [], fd_nas)
        printGolangConstStruct(gmmConstKeyValue, fd_nas)
        printGmmDecodeFunc(gmmConstKeyValue, fd_nas)
        printGmmEncodeFunc(gmmConstKeyValue, fd_nas)

        # Print GSM
        gsmMsgType = MsgTypeContent(msgDoc.tables[1])
        gsmConstKeyValue = gsmMsgType.generateMsgType()
        printGolangStruct('GsmMessage', GsmMessage, [], fd_nas)
        printGolangConstStruct(gsmConstKeyValue, fd_nas)
        printGsmDecodeFunc(gsmConstKeyValue, fd_nas)
        printGsmEncodeFunc(gsmConstKeyValue, fd_nas)


def printGolangStructGetFunc(funcDef : dict, bufName : str, fd : any) -> None :
    funcDef['returnName'] = funcDef['funcName']
    funcDef['parameterName'] = stringFirstCharToLittle(funcDef['funcName'])
    if funcDef.get('isArray', 0) == 1 :
        if funcDef['Row'][0] == funcDef['Row'][-1] :
            arrayRange = f"[{funcDef['Row'][0]}]"
        else :
            arrayRange = f"[{funcDef['Row'][0]}:{funcDef['Row'][-1] + 1}]"
    else : 
        arrayRange = ''

    if funcDef['Len']  == -1 :
        funcDef['returnType'] = '[]uint8'
        parameterName = stringFirstCharToLittle(funcDef['funcName'])
        if funcDef['Row'][0] == 0:
            content = [f"{parameterName} = make([]uint8, len(a.Buffer))", f"copy({parameterName}, a.Buffer)", f"return {parameterName}" ]
        else :
            content = [f"{parameterName} = make([]uint8, len(a.Buffer ) - {funcDef['Row'][0]})", f"copy({parameterName}, a.Buffer[{funcDef['Row'][0]}:])", f"return {parameterName}" ]
    elif bufName == 'Len' :
        funcDef['returnType'] = f"uint{funcDef['Len']}"
        content = [f"return a.{bufName}"]
    elif funcDef['Len'] % 8 != 0 and funcDef['Len'] > 8 :
        funcDef['returnType'] = f"uint{8 * (funcDef['Len'] // 8 + 1)}"
        bitSave, nowLen = f"(uint16(a.{bufName}[{funcDef['Row'][0]}]) << {funcDef['Len'] - 8}", funcDef['Len'] - 8
        for cnt in range(1, funcDef['Len'] // 8 + 1) : 
            if nowLen >= 8 :
                bitSave += f" + uint16(a.{bufName}[{funcDef['Row'][0] + cnt}]) << {nowLen}"
            elif nowLen > 0 and nowLen < 8 : 
                bitSave += f" + uint16((a.{bufName}[{funcDef['Row'][0] + cnt}]) & GetBitMask(8, {nowLen}))>> {8-nowLen})"
            nowLen -= 8
        content = [f"return {bitSave}"]
        print(f"[Warning] Get{funcDef['funcName']} = {funcDef['Len']} Length exceeds 8 and can't be mod by 8")
    elif funcDef['Len'] > 8 :
        funcDef['returnType'] = f"[{funcDef['Len'] // 8}]uint8"
        content = []
        content.append(f"copy({funcDef['parameterName']}[:], a.{bufName}{arrayRange})")
        content.append(f"return {funcDef['parameterName']}")
    else :
        funcDef['returnType'] = 'uint8' if funcDef['Len'] <= 8 else f"[{funcDef['Len'] // 8}]uint8"
        if (funcDef['sBit'] - funcDef['Len']) == 0 :
            bitMask = '' if funcDef['Len'] == 8 else f" & GetBitMask({funcDef['sBit']}, {funcDef['sBit'] - funcDef['Len']})"
        else :
            bitMask = '' if funcDef['Len'] == 8 else f" & GetBitMask({funcDef['sBit']}, {funcDef['sBit'] - funcDef['Len']}) >> ({funcDef['sBit'] - funcDef['Len']}) "
        content = [f"return a.{bufName}{arrayRange}{bitMask}"]

    funcDef['funcName'] = 'Get' + funcDef['funcName']
    printGolangStructFunc(funcDef, content, fd)
    funcDef['funcName'] = funcDef['funcName'][3:]
    funcDef.pop('returnType', None)
    funcDef.pop('returnName', None)

def printGolangStructSetFunc(funcDef : dict, bufName : str, fd : any) -> None :
    funcDef['parameterName'] = stringFirstCharToLittle(funcDef['funcName'])
    if funcDef.get('isArray', 0) == 1 :
        if funcDef['Row'][0] == funcDef['Row'][-1] :
            arrayRange = f"[{funcDef['Row'][0]}]"
        else :
            arrayRange = f"[{funcDef['Row'][0]}:{funcDef['Row'][-1] + 1}]"
    else : 
        arrayRange = ''

    if funcDef['Len']  == -1 : 
        funcDef['parameterType'] = '[]uint8'
        if funcDef['Row'][0] == 0 :
            content = [f"copy(a.Buffer, {funcDef['parameterName']})"]
        else :
            content = [f"copy(a.Buffer[{funcDef['Row'][0]}:], {funcDef['parameterName']})"]
    
    elif bufName == 'Len' :
        funcDef['parameterType'] = f"uint{funcDef['Len']}"
        content = [f"a.Len = {funcDef['parameterName']}"]
        if funcDef.get('needToMakeSlice', None) == 1 :
            content.append(f'a.Buffer = make([]uint8, a.Len)')
    elif funcDef['Len'] % 8 != 0 and funcDef['Len'] > 8 :
        funcDef['parameterType'] = f"uint{8 * (funcDef['Len'] // 8 + 1)}"
        content, nowLen = [], funcDef['Len']
        for cnt in range(0, funcDef['Len'] // 8 + 1) : 
            if nowLen > 8 :
                content.append(f"a.{bufName}[{funcDef['Row'][0] + cnt}] = uint8(({funcDef['parameterName']}) >> {nowLen - 8}) & 255")
            elif nowLen > 0 and nowLen < 8 :
                content.append(f"a.{bufName}[{funcDef['Row'][0] + cnt}] = a.{bufName}[{funcDef['Row'][0] + cnt}] & GetBitMask({8 - nowLen}, {8 - nowLen}) + uint8({funcDef['parameterName']} & {2 ** nowLen - 1}) << {8 - nowLen}")
            nowLen -= 8
        print(f"[Warning] Set{funcDef['funcName']} = {funcDef['Len']} Length exceeds 8 and can't be mod by 8")
    elif funcDef['Len'] > 8 :
        funcDef['parameterType'] = f"[{funcDef['Len'] // 8}]uint8"
        content = [f"copy(a.{bufName}{arrayRange}, {funcDef['parameterName']}[:])"]
    else :
        funcDef['parameterType'] = 'uint8' if funcDef['Len'] <= 8 else f"[{funcDef['Len'] // 8}]uint8"
        nonMotifyBitMask = 2 ** 8 - 2 ** funcDef['sBit'] + 2 ** (funcDef['sBit'] - funcDef['Len']) - 1
        nonMotify = f"" if nonMotifyBitMask == 0 else f"(a.{bufName}{arrayRange} & {nonMotifyBitMask}) + "
        
        motifyBitMask, motifyBitShift = 2 ** (funcDef['Len']) - 1, funcDef['sBit'] - funcDef['Len']
        motifyBit = funcDef['parameterName'] if motifyBitMask == 255 else f"({funcDef['parameterName']} & {motifyBitMask})"
        motify = f"{motifyBit}" if motifyBitShift == 0 else f"({motifyBit} << {motifyBitShift})"

        content = [f"a.{bufName}{arrayRange} = {nonMotify}{motify}"]

    funcDef['funcName'] = 'Set' + funcDef['funcName']
    printGolangStructFunc(funcDef, content, fd)
    funcDef['funcName']  = funcDef['funcName'][3:]
    funcDef.pop('parameterType', None)
    funcDef.pop('parameterName', None)

def generateNasCommonType() -> dict :
    ieStructToBufNameDict = dict()
    isComment = re.compile('^[ \t]*//')
    isTypeDef = re.compile('^[ \t]*type')
    isArray = re.compile('.*\[[0-9]+\][ \t]*uint8')
    
    with open('nas_type_manual.go', 'r') as fin :
        ieCollection, structContent, structFunc = dict(), dict(), dict()
        nasTypeManualRawData, structFunc['comment'] = list(), list()
        for line in fin :
            if len(line.strip()) == 0 : 
                continue
            nasTypeManualRawData.append(line)
            if isComment.match(line) != None :
                line = line.strip()[2:].strip().split() 
                ieName, ieInfo = line[0], ' '.join(line[1:])
                if 'Row' in ieInfo :
                    ieCollection[ieName] = ieInfo
                else :
                    structFunc['comment'].append(f'{ieName} {ieInfo}')
            elif isTypeDef.match(line) != None :
                line = line.split()[1]
                if structFunc.get('structName', '') != line :
                    structFunc['structName'] = line
                    fout = open(f'nasType/NAS_{line}.go', mode='w+')
                    printGolangPackage("nasType", fout)
            elif '}' in line :
                for rawData in nasTypeManualRawData : 
                    fout.write(rawData)
                print('', file=fout)

                # Like constructor
                newFuncDict, newFuncContent = dict(), list()
                newFuncDict['funcName'] = f"New{structFunc['structName']}"
                newFuncDict['returnName'], newFuncDict['returnType'] = structFunc['structName'], f"*{structFunc['structName']}"
                newFuncContent.append(f"{stringFirstCharToLittle(structFunc['structName'])} = &{structFunc['structName']}" + '{}')
                if 'Iei' in structContent or 'Iei' in ieCollection :
                    newFuncDict['parameterName'] = 'iei'
                    newFuncDict['parameterType'] = 'uint8'
                    newFuncContent.append(f"{stringFirstCharToLittle(structFunc['structName'])}.SetIei({newFuncDict['parameterName']})")
                newFuncContent.append(f"return {stringFirstCharToLittle(structFunc['structName'])}")
                printGolangStructFunc(newFuncDict, newFuncContent, fout)

                if 'Iei' in structContent and 'Iei' not in ieCollection :
                    structFunc['funcName'], structFunc['Len'] = 'Iei', 8
                    structFunc['Row'], structFunc['sBit'], bufName = [0, 0], 8, 'Iei'
                    structFunc['comment'].append(f'Iei Row, sBit, len = [], 8, 8')                    
                    printGolangStructGetFunc(structFunc, bufName, fout)
                    printGolangStructSetFunc(structFunc, bufName, fout)
                    structFunc['comment'].pop()

                if 'Len' in structContent and 'Len' not in ieCollection :
                    structFunc['funcName'], structFunc['Len'] = 'Len', int(structContent['Len'].split('uint')[-1])
                    structFunc['Row'], structFunc['sBit'], bufName  = [0, 0], 8, 'Len'
                    structFunc['comment'].append(f"Len Row, sBit, len = [], 8, {structFunc['Len']}")
                    printGolangStructGetFunc(structFunc, bufName, fout)
                    if structContent.get('Buffer', None) != None :
                        structFunc['needToMakeSlice'] = 1
                    printGolangStructSetFunc(structFunc, bufName, fout)
                    structFunc['comment'].pop()
                    structFunc.pop('needToMakeSlice', None)

                if 'Octet' in structContent :
                    structFunc['isArray'] = 1 if isArray.match(structContent['Octet']) != None else 0
                    bufName = 'Octet'
                else : 
                    structFunc['isArray'] = 1 if structFunc['Len'] != -1 else 0
                    bufName = 'Buffer'
                ieStructToBufNameDict[structFunc['structName']] = bufName

                global INF
                INF = -1
                for ieName, ieInfo in ieCollection.items() :
                    structFunc['funcName']  = ieName
                    structFunc['comment'].append(f'{ieName} {ieInfo}')
                    
                    ieInfo = re.sub('len', 'Len', ieInfo)
                    exec(ieInfo, globals(), structFunc)
                    structFunc['funcName']  = ieName
                    printGolangStructGetFunc(structFunc, bufName, fout)
                    printGolangStructSetFunc(structFunc, bufName, fout)
                    structFunc['comment'].pop()
                
                #print(f"[Info] Generate {structFunc['structName']} Get/Set Function Finish")
                structContent.clear()
                ieCollection.clear()
                structFunc['comment'].clear()
                nasTypeManualRawData.clear()
                structFunc['isArray'] = 0
            else :
                line = line.strip()
                if len(line) != 0 :
                    line = line.split()
                    structContent[line[0]] = line[1]
    return ieStructToBufNameDict


if __name__ == '__main__' : 
    
    ieStructToBufNameDict = generateNasCommonType()
    generateMessageAndAPI(ieStructToBufNameDict)
    
