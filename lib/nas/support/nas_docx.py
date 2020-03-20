#!/usr/bin/env python3

from docx import Document
import re

def stringToCapitalWithoutSpace(origin : str) -> str:
    ret = [s.strip() if s.isupper() else s.strip().title() for s in origin.split()]
    if ret[0][0].isdigit() == True :
        if len(ret) == 1 :
            if ret[0][0:3] == '5GS' :
                return ret[0][3:] + '5GS'
            elif ret[0][0:2] == '5G' :
                return ret[0][2:] + '5G'
            else :
                print('[stringToCapitalWithoutSpace] Special Case Need To Do')
        ret.append(ret[0])
        return ''.join(ret[1:])
    return ''.join(ret)

def printGolangStruct(structName : str, content : list(), comment : list, fd : any) -> None: 
    for line in comment :
        print(f'// {line}', file=fd)
    print('type', structName, 'struct {', file=fd)
    for line in content :
        print(f'    {line}', file=fd)
    print('}\n', file=fd)

def printGolangStructFunc(funcDef : dict, content : list(), fd : any) -> None :
    print(f"// {funcDef.get('funcName', '')} {funcDef.get('comment', '')}", file=fd)
    print(f"func (a *{funcDef.get('structName', '')}) {funcDef.get('funcName', '')}() ({funcDef.get('funcName', '')} {funcDef.get('returnType', '')})", '{', file=fd)
    for line in content : 
        print(f'    {line}', file=fd)
    print('}\n', file=fd)

class MessageContent() :
    def __init__(self, name = '', section = '') :
        self.IE = list()
        self.name = name.strip()
        self.section = section.strip()

    def getInfo(self, table : list) -> None :
        titleList = list()
        for cells in table.rows[0].cells :
            titleList.append(cells.text)

        for row in table.rows[1:] :
            tmpDict = dict()
            for title, cell in zip(titleList, row.cells) :
                tmpDict[title] = cell.text.strip()
            self.IE.append(tmpDict)

    def generateStruct(self) -> list:
        ieCollection = list()
        typeName = stringToCapitalWithoutSpace(re.sub('[\'\-\"]', '', self.name)).split('(')[0]
        contentList, commentList, varName, isHalf = list(), list(), list(), 0
        for ie in self.IE :
            varName.append(stringToCapitalWithoutSpace(re.sub('[\'\-\"]', '', ie['Information Element'])).split('(')[0])
            commentList.append(ie['Type/Reference'].split()[-1])
            if ie['Length'] != '1/2' and isHalf == 1 :
                print('[Not a commplete byte]')
            if ie['Length'] == '1/2' and isHalf == 0 :
                isHalf = 1
                continue

            ieName = 'And'.join(varName)
            if ie['Presence'] == 'M' :
                contentList.append(ieName)
            else :
                contentList.append(f'*{ieName}')

            ieCollection.append([ieName, ie['Format'], ie['Length'], ' '.join(commentList)])
            commentList, varName, isHalf = list(), list(), 0
        
        if isHalf == 1 :
            print('[Not a commplete byte]')

        printGolangStruct(typeName, contentList, commentList, fout)
        return ieCollection

    def print(self) -> None :
        print(f'[{self.name}]')
        for ie in self.IE :
            print(ie)


class IEContent() : 
    def __init__(self, name : str, format : str, length : str, comment : str) :
        self.name = name.strip()
        self.format = format.strip()
        self.length = length.strip()
        self.comment = comment.strip()

    def generateStruct(self, fd : any) -> None :
        comment = [f'{self.name} {self.comment}']
        if self.length == '1/2' : 
            content = ['Octet uint8']
        elif self.format == 'V' :
            if self.length == '1' and self.comment == '9.7' :
                content = ['Octet uint8']
                comment.append('MessageType Row, sBit, len = [0, 0], 8 , 8')
            else :
                try : 
                    octetLen = '' if int(self.length) == 1 else f'[{int(self.length)}]'
                    content = ['Iei uint8', f'Octet {octetLen}uint8']
                except :
                    print(f'[{self.name}] Length Error')
                    content = []
        else : 
            content = []
        printGolangStruct(self.name, content, comment, fd)

    def generateGetFunc(self, fd : any) -> None :
        funcDef = {'structName' : self.name}
        if self.length == '1/2' : 
            funcDef['returnType'] = 'unit8'

            funcDef['funcName'] = f"Get{self.name.split('And')[0]}"
            funcDef['comment'] = self.comment.split()[0]
            content = [f'return a.{self.name} & GetBitMask(8, 4)']
            printGolangStructFunc(funcDef, content, fd)

            funcDef['funcName'] = f"Get{self.name.split('And')[1]}"
            funcDef['comment'] = self.comment.split()[1]
            content = [f'return a.{self.name} & GetBitMask(4, 0)']
            printGolangStructFunc(funcDef, content, fd)

    def print(self) :
        print(f'// [{self.name}] format : {self.format} length : {self.length} comment : {self.comment}')


contentDoc = Document('content_type.docx')

messages = list()
ieList, ieCollection = list(), set()

for line in contentDoc.paragraphs :
    if line.style.name == 'Heading 3' :
        h3TitleList = line.text.strip().split()
        messageName, comment = ' '.join(h3TitleList[1:]), h3TitleList[0]
        messages.append(MessageContent(messageName, comment))

with open('nas_message.go', 'w') as fout :
    for idx, table in enumerate(contentDoc.tables) : 
        messages[idx].getInfo(table)
        ies = messages[idx].generateStruct()
        for ie in ies : 
            if ie[0] in ieCollection :
                continue
            ieList.append(IEContent(*ie))
            ieCollection.add(ie[0])


with open('nas_type.go', 'w') as fd_type :
    with open('nas_get_func.go', 'w') as fd_func :
        for ie in ieList :
            ie.generateStruct(fd_type)
            ie.generateGetFunc(fd_func)
