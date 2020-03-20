#!/usr/bin/env python3

# Please save the doc as docx before delete useless table.
# Check all of table are complete. There are problems if rows are not align in table.
# All tables are saved in variable "tables" using structure "list".
# Rows for each table use structure "dict" and save in variable "tables[index]".

import os, re
from docx import Document

ieNameMapping = {
    'PFD': 'PFDContext',
    'PDRID': 'PacketDetectionRuleID',
    'SxSRRspFlags': 'PFCPSRRspFlags'
}

def formatString(inputStr : str) -> str :
    inputStr = re.sub(r"['/\"]", '', inputStr)
    outputStrList = [s[0].upper() + s[1:] for s in re.sub(r'[-() ]+', ' ', inputStr).split()]
    return ''.join(outputStrList)

class FileOutput() :
    def __init__(self, fileName : str) :
        self.fd = open(fileName, 'w')
        self.tab, self.tabstop = 0, 4

    def indent(self, num : int) :
        self.tab += num
        self.tab = 0 if self.tab < 0 else self.tab

    def indented(self, contentList : list) :
        self.indent(+1)
        for content in contentList :
            self.fprint(content)
        self.indent(-1)

    def fprint(self, content : str) :
        print(' ' * self.tabstop * self.tab, content, sep='', file=self.fd)

class TableParser() :
    def __init__(self, fileName : str) :
        self.document = Document(fileName)
        self.tables = []
        self.parse()

    def parse(self) :
        for idx, table in enumerate(self.document.tables) :
            gotTitle, titlePointer = 0, None
            for row in table.rows :
                try :
                    if 'Information elements'.lower() in [cell.text.lower() for cell in row.cells] :
                        if gotTitle == 0 :
                            self.tables.append(list())
                        titlePointer, gotTitle = row, 1
                    elif gotTitle == 1 :
                        content, isNote = dict(), 0
                        for title, context in zip(self.yieldTitleFromDocx(titlePointer), row.cells) :
                            if context._tc.right - context._tc.left >= 8 :
                                isNote = 1
                                break
                            content[title] = context.text
                        if isNote == 0 :
                            self.tables[-1].append(content)
                        
                except :
                    print(f'[Error] The {idx} table is dirty')
                    break

    def yieldTitleFromDocx(self, tableRowPtr) :
        for cell in tableRowPtr.cells :
            yield cell.text

    def printTableByIndex(self, idxOfTable) :
        try :
            for content in self.tables[idxOfTable] :
                print(content)
        except :
            print('[Warning] Index out of bound')


if __name__ == '__main__' :

    doc29244_812_1 = TableParser('29244-f30-ch8.1.2-1.docx')
    ieMappingValue = dict()
    for row in doc29244_812_1.tables[0][:-1] :
        ieName, ieVal = formatString(row['Information elements']), row['IE Type value\n(Decimal)']
        if ieMappingValue.get(ieName, None) == None :
            ieMappingValue[ieName] = ieVal
        else :
            print(f'[Warning] {ieName} is duplicate')

    specialCase = set()
    specialCase.update(['UpdateBAR', 'UsageReport'])

    # There have 67 table in chapter 7, but the first one will not be used
    docxChapter7Name = '29244-f30-ch7-fixed-table.docx'
    doc29244_7_para = Document(docxChapter7Name)
#    tableName = re.compile(r'Table 7.*: (Information Elements in [an ]{0,3})?(.+(?= IE within ))?(.+)')
    tableName = re.compile(r'Table 7.*: (Information Elements in [an ]{0,3}|(.+)( IE within ))?(.+)')
    chapter7TitleList = []
    for line in doc29244_7_para.paragraphs :
        afterMatch = tableName.match(line.text)
        if afterMatch :
            ieName = afterMatch.group(2) if afterMatch.group(2) else afterMatch.group(4)
            if formatString(ieName) in specialCase :
                ieName += afterMatch.group(4)
            chapter7TitleList.append(ieName)
            # print(afterMatch.group(2)) if afterMatch.group(2) else print(afterMatch.group(3))

    doc29244_7 = TableParser(docxChapter7Name)
    chapter7UsedIESet = set()
    for tableName in chapter7TitleList[1:] :
        chapter7UsedIESet.add(formatString(tableName))

    PFCPMessageTypeFd = FileOutput('pfcpMessage.go')
    PFCPMessageTypeFd.fprint('package pfcpMessage\n')
    PFCPMessageTypeFd.fprint('import "gofree5gc/lib/pfcp/pfcpType"\n')
    for tableName, table in zip(chapter7TitleList[1:], doc29244_7.tables) :
        tableName = formatString(tableName)
        PFCPMessageTypeFd.fprint(f'type {tableName} struct ' +'{')
        PFCPMessageTypeFd.indent(+1)
        # Skip PFCP Session Deletion Request
        for ie in table :
            if tableName == 'PFCPSessionDeletionRequest':
                break
            try :
                ieName = formatString(ie['Information elements'])
            except :
                ieName = 'NoIEName'
                print(f'[warning] No IE name in {tableName}')
            try :
                ieType = formatString(ie['IE Type'])
            except :
                ieType = 'NoIEType'
                print(f'[warning] No IE {ieName} type in {tableName}')
            try :
                if ieNameMapping.get(ieType) :
                    ieType = ieNameMapping[ieType]
                elif ieType in specialCase : 
                    ieType += tableName
                ieTLV = f' `tlv:"{ieMappingValue[ieType]}"`'
            except :
                ieTLV = ''
            if len(ieType) > 0 and ieType not in chapter7UsedIESet :
                PFCPMessageTypeFd.fprint(f'{ieName} *pfcpType.{ieType}{ieTLV}')
            else :
                PFCPMessageTypeFd.fprint(f'{ieName} *{ieType}{ieTLV}')

        PFCPMessageTypeFd.indent(-1)
        PFCPMessageTypeFd.fprint('}\n')


    os.makedirs('pfcpType', 0o755, exist_ok=True)
    for ieName, ieVal in ieMappingValue.items() :
        if ieName not in chapter7UsedIESet:
            PFCPBaseTypeFd = FileOutput(f'pfcpType/{ieName}.go')
            PFCPBaseTypeFd.fprint('package pfcpType\n')
            PFCPBaseTypeFd.fprint(f'type {ieName} struct ' +'{')
            PFCPBaseTypeFd.indented([f'{ieName}value []byte'])
            PFCPBaseTypeFd.fprint('}\n')
